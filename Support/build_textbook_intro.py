"""Generate the Chinese review-oriented introduction Word document."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).parent / "textbook_intro_review.docx"

CN_FONT = "宋体"
EN_FONT = "Times New Roman"


def set_run_font(run, size_pt, bold=False, color=None):
    run.font.name = EN_FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 22, bold=True)
    p.paragraph_format.space_after = Pt(12)


def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 14, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, 12)


def add_image_placeholder(doc, caption):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Cm(14)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[此处插入{caption}]")
    set_run_font(run, 11, color=RGBColor(0x88, 0x88, 0x88))

    # Approximate a 5cm-tall placeholder by padding the cell.
    for _ in range(3):
        extra = cell.add_paragraph()
        extra.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(extra.add_run(""), 11)

    # Caption under the box.
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, 10.5)
    cap.paragraph_format.space_after = Pt(8)


def main():
    doc = Document()

    # A4 with tight margins to keep everything tidy.
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)

    add_title(doc, "《新工科数学分析》讲义简介")

    # Subtitle / affiliation
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(sub.add_run("首都师范大学  ·  面向人工智能时代的数学分析教材"), 11)
    sub.paragraph_format.space_after = Pt(10)

    # Section 1 —— 定位
    add_heading(doc, "一、教材定位")
    add_body(
        doc,
        "《新工科数学分析》是为首都师范大学新工科类专业量身打造的数学分析讲义,"
        "在保持分析学严谨体系的同时,系统地把极限、微分、积分与向量分析的经典工具,"
        "重新组织为支撑人工智能、数据科学与工程建模的数学基础,力图为新工科学生"
        "提供一条"
        "\u201c经典分析—计算实现—智能应用\u201d贯通的学习路径。"
    )

    # Section 2 —— AI 落脚
    add_heading(doc, "二、内容以人工智能与工程建模为落脚点")
    add_body(
        doc,
        "讲义全书十五章覆盖一元、多元、向量分析与积分理论,每章均设置与人工智能、"
        "动力系统、信号处理等前沿场景密切相关的典型案例。例如:在微分方程一章,"
        "以洛特卡-沃尔泰拉方程、弹簧振子、行星运动阐释非线性建模;在多元微分与向量场"
        "部分,引入梯度下降与神经网络训练的几何解释;在积分变换章节,借助 FFT 讲解"
        "信号与数据的频域结构。这样的安排,使学生不仅能掌握定理与公式,更能在"
        "真实的智能计算任务中识别并使用数学工具,符合新工科人才培养对"
        "\u201c懂原理、会建模、能落地\u201d的总体要求。"
    )
    add_image_placeholder(doc, "图 1  网页版章节渲染效果(公式、插图、讲解一体化排版)")

    # Section 3 —— 多平台
    add_heading(doc, "三、网页、移动端与三维交互一体的多平台阅读体验")
    add_body(
        doc,
        "讲义突破传统纸质教材的单向呈现,提供网页版、Android 移动端 App 与可交互图形"
        "三位一体的阅读形态。网页版基于 VLOOK 主题呈现,数学排版清晰、检索便捷;"
        "章节中嵌入十余个 Three.js 三维场景,读者可直接旋转、缩放梯度场、双重积分区域、"
        "雅可比变换、斯托克斯曲面等关键几何对象;Plotly 动态图则支持参数调节与数据交互。"
        "与此同时,配套的 Android App 让学生在手机上即可阅读全部内容、播放三维演示、"
        "查看例题与习题,真正做到\u201c随身课堂\u201d,显著提升自主学习与课后复盘的效率。"
    )
    add_image_placeholder(doc, "图 2  Three.js 三维交互插图示例(如双重积分或梯度场)")
    add_image_placeholder(doc, "图 3  Android 移动端 App 阅读界面")

    # Section 4 —— 可运行代码
    add_heading(doc, "四、可运行的教学代码,公式—代码—可视化三位一体")
    add_body(
        doc,
        "讲义配套 40 余个 Python 教学脚本,覆盖极限可视化、导数近似、凸函数判别、"
        "自由落体与弹簧振动、行星轨道、神经网络前向传播、傅里叶变换等。每个脚本均可"
        "独立运行并生成对应插图,学生在课堂或课下可一边对照公式,一边修改参数观察结果,"
        "从而在动手实验中加深对抽象概念的理解。"
    )
    add_image_placeholder(doc, "图 4  Python 建模脚本可视化输出(如行星运动)")

    # Section 5 —— 习题与工程化
    add_heading(doc, "五、完整习题体系与工程化构建管线")
    add_body(
        doc,
        "讲义已编制 217 道配套习题,按章节组织、难度分层,并由独立的构建管线统一生成"
        "网页、App 与打印版本,实现\u201c一次编写、多端发布\u201d。工程化的写作与发布流程"
        "保证了教材的长期可维护性,也为后续内容迭代、分册出版与配套资源扩展奠定了"
        "稳固的技术基础。"
    )

    doc.save(OUT)
    print(f"written: {OUT}")


if __name__ == "__main__":
    main()
